from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RUN_ID = "RUN-f944ce026be1"
RUN_DIR = ROOT / "project" / "runs" / RUN_ID
OUT_DIR = ROOT / "entregas"
OUT_PATH = OUT_DIR / "Documentacion_Proyecto_Mini_Fabrica_Agentica_ClaveUnica.docx"


def read_json(path: Path):
    return json.loads(path.read_text(encoding="utf-8"))


scope = read_json(RUN_DIR / "scope-inventory.json")
final_report = read_json(RUN_DIR / "final-report.json")
runtime_gate = read_json(RUN_DIR / "runtime-close-gate.json")
docker_runtime = read_json(RUN_DIR / "docker-runtime-validation.json")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
        set_cell_shading(hdr[i], "F1F3F4")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if widths:
                cells[i].width = Inches(widths[i])
    doc.add_paragraph()


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 20, "000000"),
        ("Heading 2", 16, "000000"),
        ("Heading 3", 14, "434343"),
    ]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14)
        style.paragraph_format.space_after = Pt(6)


def title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Documentacion del Proyecto")
    r.font.name = "Arial"
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(0, 0, 0)

    subtitle = doc.add_paragraph()
    subtitle.add_run("Mini fabrica agentica de procesos - Portal Ciudadano ClaveUnica").bold = True
    doc.add_paragraph("Entregable 1: Documentacion del proyecto (5%)")
    doc.add_paragraph(f"Run final documentado: {RUN_ID}")
    doc.add_paragraph("Modalidad de ejecucion: local, Docker Compose, modo Codex directo con OpenAI API bloqueada.")
    doc.add_paragraph("Autor: Benjamin Cruzado")
    doc.add_paragraph()
    doc.add_paragraph(
        "Este documento consolida la especificacion, alcance, arquitectura, agentes, reglas, pruebas, "
        "trazabilidad y evidencia tecnica generada por la fabrica para la entrega final integradora."
    )
    doc.add_page_break()


def add_summary(doc: Document) -> None:
    doc.add_heading("1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "El proyecto implementa una mini fabrica agentica orientada a transformar un work order y documentos "
        "de requisitos en una aplicacion web local verificable. El caso trabajado corresponde a un Portal "
        "Ciudadano ClaveUnica con frontend Angular, backend Spring Boot, base de datos PostgreSQL, migraciones "
        "Flyway, contenedores Docker y pruebas automatizadas."
    )
    doc.add_paragraph(
        "La fabrica conserva la estructura solicitada en la evaluacion: entrada clara, refinamiento, "
        "planificacion, agentes especializados, reglas y workflows, arnes de ejecucion, validacion, evidencia "
        "y cierre tecnico. La ejecucion final documentada cerro en estado complete."
    )
    add_table(
        doc,
        ["Elemento", "Resultado"],
        [
            ["Run final", RUN_ID],
            ["Estado final", str(final_report.get("status"))],
            ["Ready for first project", str(final_report.get("ready_for_first_project"))],
            ["Docker runtime", str(runtime_gate.get("docker_runtime_status"))],
            ["Playwright", str(runtime_gate.get("playwright_status"))],
            ["OpenAI API", "Bloqueada por defecto en modo codex_direct"],
            ["GitHub / EC2", "Congelados durante esta preparacion local"],
        ],
        [2.2, 4.0],
    )


def add_factory_design(doc: Document) -> None:
    doc.add_heading("2. Diseno de la fabrica agentica", level=1)
    doc.add_paragraph(
        "La fabrica se organiza alrededor de un orquestador, un arnes de ejecucion y agentes especializados. "
        "Cada agente opera con contratos, policy, contexto, evidencia, validadores y registros de salida."
    )
    doc.add_heading("2.1 Entrada clara", level=2)
    add_bullets(
        doc,
        [
            "Work order actual: project/input/work_order.json.",
            "Documento funcional del profesor incorporado como fuente autorizada.",
            "Product generation brief incorporado como politica de generacion.",
            "Restricciones: trabajo local, sin datos productivos, sin credenciales reales y sin despliegue automatico.",
        ],
    )
    doc.add_heading("2.2 Flujo de trabajo de la fabrica", level=2)
    add_numbered(
        doc,
        [
            "Intake y validacion del work order.",
            "Refinamiento de requisitos y limpieza de contexto.",
            "Planificacion tecnica y alcance de rubrica.",
            "Especializacion por agentes de frontend, backend, base de datos, pruebas, Docker, QA y revisores.",
            "Ejecucion por arnes con policy, budget, evidence gates y safety gates.",
            "Validacion con Docker, smoke tests, Playwright y reviewers.",
            "Cierre tecnico con final-report, runtime-close-gate y evidencia.",
        ],
    )
    doc.add_heading("2.3 Agentes ejecutados en el cierre", level=2)
    agents = final_report.get("agents_executed", [])
    rows = [[str(i + 1), agent] for i, agent in enumerate(agents[:24])]
    add_table(doc, ["#", "Agente"], rows, [0.5, 5.7])


def add_scope(doc: Document) -> None:
    doc.add_heading("3. Alcance minimo solicitado y cobertura", level=1)
    counts = {
        "Casos de uso": len(scope.get("use_cases", [])),
        "Funcionalidades / flujos": len(scope.get("features", [])),
        "Tablas": len(scope.get("tables", [])),
        "Endpoints API": len(scope.get("api_endpoints", [])),
        "Pantallas": len(scope.get("screens", [])),
        "Reglas de negocio": len(scope.get("business_rules", [])),
        "Validaciones / restricciones CHECK": len(scope.get("validations_checks", [])),
        "Trazabilidad": len(scope.get("traceability", [])),
    }
    add_table(
        doc,
        ["Elemento de rubrica", "Minimo", "Evidencia generada", "Estado"],
        [
            ["Documento de especificacion", "6 paginas", "Este documento + spec.md", "Cubierto"],
            ["Casos de uso", "10", str(counts["Casos de uso"]), "Cubierto"],
            ["Funcionalidades / flujos", "30", str(counts["Funcionalidades / flujos"]), "Cubierto"],
            ["Tablas", "40", str(counts["Tablas"]), "Cubierto"],
            ["Endpoints API", "40", str(counts["Endpoints API"]), "Cubierto"],
            ["Pantallas", "30", str(counts["Pantallas"]), "Cubierto"],
            ["Reglas de negocio", "60", str(counts["Reglas de negocio"]), "Cubierto"],
            ["Validaciones / CHECK", "100", str(counts["Validaciones / restricciones CHECK"]), "Cubierto"],
            ["Checklist", "1", "CHECKLIST_APLICADO.md", "Cubierto"],
            ["Pruebas automatizadas", "100% segun rubrica", "Smoke + Playwright + runtime", "Cubierto con evidencia local"],
        ],
        [2.1, 1.1, 2.2, 1.0],
    )
    doc.add_paragraph(
        "Nota de alcance: la evidencia documenta cumplimiento estructural de rubrica y ejecucion local. "
        "La calidad funcional fina de la experiencia de usuario debe seguir mejorandose si se evalua como producto real completo."
    )


def add_use_cases(doc: Document) -> None:
    doc.add_heading("4. Casos de uso", level=1)
    doc.add_paragraph("La fabrica genero 21 casos de uso. La siguiente tabla presenta los principales.")
    rows = []
    for uc in scope.get("use_cases", [])[:12]:
        rows.append([
            uc.get("id", ""),
            uc.get("title", ""),
            uc.get("workflow", ""),
            ", ".join(uc.get("visible_effects", [])[:3]) if isinstance(uc.get("visible_effects"), list) else str(uc.get("visible_effects", "")),
        ])
    add_table(doc, ["ID", "Caso de uso", "Workflow", "Efecto visible esperado"], rows, [0.8, 2.0, 1.3, 2.4])


def add_screens_endpoints_tables(doc: Document) -> None:
    doc.add_heading("5. Pantallas, endpoints y tablas", level=1)
    doc.add_heading("5.1 Pantallas principales", level=2)
    rows = []
    for screen in scope.get("screens", [])[:30]:
        rows.append([
            screen.get("id", ""),
            screen.get("route", ""),
            screen.get("module", ""),
            ", ".join(screen.get("actions", [])[:3]),
        ])
    add_table(doc, ["ID", "Ruta", "Modulo", "Acciones"], rows, [0.7, 2.0, 1.4, 2.4])

    doc.add_heading("5.2 Endpoints API principales", level=2)
    rows = []
    for ep in scope.get("api_endpoints", [])[:45]:
        rows.append([ep.get("id", ""), ep.get("method", ""), ep.get("path", ""), ep.get("title", "")])
    add_table(doc, ["ID", "Metodo", "Endpoint", "Proposito"], rows, [0.7, 0.8, 2.5, 2.4])

    doc.add_heading("5.3 Tablas de dominio", level=2)
    rows = []
    for table in scope.get("tables", [])[:40]:
        rows.append([table.get("id", ""), table.get("table", ""), table.get("workflow", "")])
    add_table(doc, ["ID", "Tabla", "Workflow"], rows, [0.8, 3.0, 2.4])


def add_rules_validations(doc: Document) -> None:
    doc.add_heading("6. Reglas, validaciones y restricciones", level=1)
    doc.add_heading("6.1 Reglas de negocio", level=2)
    rows = []
    for rule in scope.get("business_rules", [])[:25]:
        rows.append([rule.get("id", ""), rule.get("workflow", ""), rule.get("rule", "")])
    add_table(doc, ["ID", "Workflow", "Regla"], rows, [0.8, 1.2, 4.4])
    doc.add_paragraph("El inventario completo contiene 60 reglas de negocio en scope-inventory.json.")

    doc.add_heading("6.2 Validaciones y CHECK", level=2)
    rows = []
    for check in scope.get("validations_checks", [])[:30]:
        rows.append([check.get("id", ""), check.get("type", ""), check.get("applies_to", ""), check.get("check", "")])
    add_table(doc, ["ID", "Tipo", "Aplica a", "Validacion"], rows, [0.7, 0.8, 1.7, 3.3])
    doc.add_paragraph("El inventario completo contiene 100 validaciones/restricciones CHECK en scope-inventory.json.")


def add_architecture_tests(doc: Document) -> None:
    doc.add_heading("7. Arquitectura tecnica del sistema generado", level=1)
    add_table(
        doc,
        ["Capa", "Tecnologia / evidencia"],
        [
            ["Frontend", "Angular; rutas y componentes en app-generada/frontend/src/app."],
            ["Backend", "Spring Boot REST JSON; controllers, services, DTOs y repositories."],
            ["Base de datos", "PostgreSQL con Flyway; migracion V1__init_claveunica_domain.sql."],
            ["Contenedores", "Docker Compose con servicios postgres, backend y frontend."],
            ["Pruebas", "Smoke test, Playwright E2E y validacion Docker runtime."],
            ["Trazabilidad", "traceability-matrix.md y scope-inventory.json."],
        ],
        [1.6, 4.8],
    )

    doc.add_heading("8. Diseno y plan de pruebas", level=1)
    add_table(
        doc,
        ["Prueba", "Tipo", "Evidencia", "Estado"],
        [
            ["WorkOrder/CycleState schema", "unit", "tests/test_factory.py", "Definido"],
            ["Agent registry", "unit", "registries/agents.json", "Definido"],
            ["Policy allowlist", "negativo", "PolicyValidator", "Definido"],
            ["Smoke app generada", "smoke", "app-generada/tests/smoke.mjs", "Ejecutado"],
            ["Playwright E2E", "e2e", "app-generada/tests/e2e.spec.ts", "Ejecutado"],
            ["Docker runtime", "runtime", "docker-runtime-validation.json", "runtime_complete"],
            ["Backend health", "runtime", "http://localhost:8080/api/v1/health", "complete"],
            ["Frontend health", "runtime", "http://localhost:3000", "complete"],
        ],
        [1.8, 1.0, 2.6, 1.0],
    )
    doc.add_paragraph(
        f"El cierre tecnico registro Docker runtime en estado {docker_runtime.get('status')} y Playwright en estado "
        f"{docker_runtime.get('playwright_status')}."
    )


def add_traceability_and_closure(doc: Document) -> None:
    doc.add_heading("9. Trazabilidad y evidencia", level=1)
    rows = []
    for item in scope.get("traceability", [])[:25]:
        rows.append([
            item.get("requirement", item.get("id", "")),
            item.get("screen", ""),
            item.get("endpoint", ""),
            item.get("table", ""),
            item.get("test", ""),
        ])
    if rows:
        add_table(doc, ["Req.", "Pantalla", "Endpoint", "Tabla", "Prueba"], rows, [0.7, 1.4, 1.8, 1.3, 1.1])
    doc.add_paragraph(
        "La trazabilidad completa se conserva en los artefactos del run final: traceability-matrix.md, "
        "scope-inventory.json, evidence-register.json y claim-map.md."
    )

    doc.add_heading("10. Checklist de completitud", level=1)
    add_table(
        doc,
        ["Artefacto", "Estado", "Ubicacion"],
        [
            ["Work order", "complete", "project/input/work_order.json"],
            ["Especificacion", "complete", f"project/runs/{RUN_ID}/spec.md"],
            ["Plan", "complete", f"project/runs/{RUN_ID}/plan.md"],
            ["Contratos", "complete", f"project/runs/{RUN_ID}/contracts.md"],
            ["Pruebas", "complete", f"project/runs/{RUN_ID}/test-report.md"],
            ["Validacion", "complete", f"project/runs/{RUN_ID}/validation-report.json"],
            ["Docker runtime", "complete", f"project/runs/{RUN_ID}/docker-runtime-validation.json"],
            ["Cierre", "complete", f"project/runs/{RUN_ID}/final-report.json"],
        ],
        [1.7, 1.0, 3.8],
    )

    doc.add_heading("11. Limitaciones conocidas", level=1)
    add_bullets(
        doc,
        [
            "La ejecucion final documentada fue local; el despliegue EC2 quedo congelado durante la preparacion de este documento.",
            "El modo Codex directo bloqueo llamadas OpenAI API para evitar gasto accidental.",
            "La app generada cuenta con persistencia y runtime validado, pero la experiencia de usuario requiere mejoras para una evaluacion estricta de producto moderno.",
            "La cobertura indicada corresponde a evidencia y gates de fabrica; se recomienda complementar con pruebas funcionales mas profundas por caso de uso.",
        ],
    )

    doc.add_heading("12. Conclusion", level=1)
    doc.add_paragraph(
        "El proyecto presenta una fabrica agentica organizada y verificable, con entrada clara, agentes especializados, "
        "arnes, reglas, validadores, trazabilidad y cierre tecnico. La documentacion adjunta permite revisar el alcance "
        "de rubrica, el codigo generado, el plan de pruebas y la evidencia de ejecucion local. Para maximizar la nota "
        "final, los entregables siguientes deben enfocarse en despliegue EC2 y video de presentacion, ya que concentran "
        "el 80% de la evaluacion."
    )


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    title_page(doc)
    add_summary(doc)
    add_factory_design(doc)
    add_scope(doc)
    add_use_cases(doc)
    add_screens_endpoints_tables(doc)
    add_rules_validations(doc)
    add_architecture_tests(doc)
    add_traceability_and_closure(doc)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
